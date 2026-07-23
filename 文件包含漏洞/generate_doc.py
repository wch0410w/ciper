"""
生成文件包含漏洞发现与修复报告 Word 文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc, code_text):
    """添加代码块（灰色背景、等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)

    # 添加灰色背景
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)

    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def generate_report():
    doc = Document()

    # ========== 页面设置 ==========
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== 标题样式 ==========
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ========== 封面标题 ==========
    doc.add_paragraph()  # 空行
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Web应用文件包含漏洞')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('发现与修复报告')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x4A, 0x6E, 0xB5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    doc.add_paragraph()

    # 封面信息
    info_items = [
        ('项目名称', 'Flask Web 应用（文件上传漏洞演示系统）'),
        ('文档版本', 'V1.0'),
        ('编制日期', datetime.date.today().strftime('%Y年%m月%d日')),
        ('漏洞类型', '文件包含漏洞（File Inclusion / Path Traversal）'),
        ('严重等级', '🔴 高危（High）'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ========== 目录页 ==========
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('目  录')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    toc_items = [
        '一、漏洞概述',
        '二、漏洞发现过程',
        '    2.1 代码审计发现',
        '    2.2 漏洞验证测试',
        '    2.3 攻击场景分析',
        '三、漏洞成因分析',
        '四、漏洞风险评估',
        '五、漏洞修复方案',
        '六、修复验证',
        '七、安全建议',
        '附录：相关代码片段',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ==========================================
    # 一、漏洞概述
    # ==========================================
    h1 = doc.add_heading('一、漏洞概述', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph(
        '文件包含漏洞（File Inclusion Vulnerability）是指Web应用程序在包含文件时，'
        '未对用户输入的文件路径进行充分过滤和校验，导致攻击者可以通过构造恶意的文件路径，'
        '读取服务器上的任意文件（本地文件包含，LFI），甚至远程包含恶意文件（远程文件包含，RFI）。'
    )

    doc.add_paragraph(
        '在本项目中，Flask Web应用的 /page 路由存在典型的本地文件包含（LFI）漏洞。'
        '攻击者可以在未经授权的情况下，通过路径遍历（Path Traversal）技术，'
        '读取服务器上的任意文件，包括但不限于：'
    )

    vuln_list = [
        '应用程序源代码（如 app.py），泄露业务逻辑和数据库结构；',
        '数据库文件（如 data/users.db），直接获取所有用户的账号密码；',
        '系统配置文件（如 /etc/passwd、/etc/hosts），获取服务器敏感信息；',
        '日志文件（如访问日志、错误日志），获取请求记录和调试信息。',
    ]
    for item in vuln_list:
        p = doc.add_paragraph(item, style='List Bullet')

    # ==========================================
    # 二、漏洞发现过程
    # ==========================================
    h1 = doc.add_heading('二、漏洞发现过程', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # 2.1
    h2 = doc.add_heading('2.1 代码审计发现', level=2)

    doc.add_paragraph(
        '通过对应用源代码进行安全审计，在 app.py 文件的 /page 路由处理函数中发现了文件包含漏洞。'
        '该路由的 URL 为 "/page?name=xxx"，功能是根据用户传入的 name 参数加载 pages 目录下的对应 HTML 页面。'
    )

    doc.add_paragraph('核心漏洞代码如下：')

    add_code_block(doc, '''@app.route("/page")
def page():
    name = request.args.get("name", "")        # 直接获取用户输入，未做任何过滤
    page_content = None
    if name:
        filepath = os.path.join("pages", name) # ★ 直接拼接用户输入到文件路径
        if os.path.exists(filepath):
            with open(filepath, encoding="utf-8") as f:
                page_content = f.read()         # ★ 读取并返回文件内容
        else:
            filepath = os.path.join("pages", name + ".html")
            if os.path.exists(filepath):
                with open(filepath, encoding="utf-8") as f:
                    page_content = f.read()
            else:
                page_content = "页面不存在"
    ...
    return render_template("index.html", ..., page_content=page_content)''')

    doc.add_paragraph('审计发现的安全问题：')

    findings = [
        ('未校验用户输入', 'name 参数直接取自 request.args.get("name")，未进行任何白名单校验或路径过滤。'),
        ('路径拼接不安全', '使用 os.path.join("pages", name) 拼接路径，但 os.path.join 并不阻止路径遍历（如传入 "../" 可跳出 pages 目录）。'),
        ('文件读取无限制', '只要文件存在就读取并返回内容，未限制可读取的文件类型或目录范围。'),
        ('输出未转义', '模板中使用 {{ page_content | safe }} 渲染，直接将文件内容作为HTML输出，可能引发XSS。'),
    ]
    for title, desc in findings:
        p = doc.add_paragraph()
        run = p.add_run(f'● {title}：')
        run.font.bold = True
        p.add_run(desc)

    # 2.2
    h2 = doc.add_heading('2.2 漏洞验证测试', level=2)

    doc.add_paragraph(
        '为验证漏洞的真实性和危害程度，进行了以下安全测试：'
    )

    doc.add_paragraph('测试1：正常功能访问（基准测试）', style='List Number')
    doc.add_paragraph('    URL: http://localhost:5000/page?name=help')
    doc.add_paragraph('    结果：正常返回 pages/help.html 的内容，功能正常。')

    doc.add_paragraph('测试2：路径遍历读取应用源码', style='List Number')
    p = doc.add_paragraph()
    run = p.add_run('    URL: http://localhost:5000/page?name=../app.py')
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_paragraph('    结果：成功返回 app.py 的全部源代码，包括数据库结构、密钥等敏感信息。')
    doc.add_paragraph('    说明：路径 ../app.py 经 os.path.join 处理后变为 pages/../app.py，即 app.py。')

    doc.add_paragraph('测试3：路径遍历读取数据库文件', style='List Number')
    p = doc.add_paragraph()
    run = p.add_run('    URL: http://localhost:5000/page?name=../data/users.db')
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_paragraph('    结果：成功读取数据库二进制文件，可直接获取所有用户的账号密码。')

    doc.add_paragraph('测试4：路径遍历读取模板文件', style='List Number')
    p = doc.add_paragraph()
    run = p.add_run('    URL: http://localhost:5000/page?name=../templates/login.html')
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_paragraph('    结果：成功读取登录页面模板源码。')

    doc.add_paragraph('测试5：多层路径遍历（Linux环境）', style='List Number')
    p = doc.add_paragraph()
    run = p.add_run('    URL: http://localhost:5000/page?name=../../../etc/passwd')
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_paragraph('    结果：成功读取 /etc/passwd 系统文件，列出所有系统用户。')

    # 2.3
    h2 = doc.add_heading('2.3 攻击场景分析', level=2)

    doc.add_paragraph('场景一：源代码泄露', style='List Number')
    doc.add_paragraph(
        '攻击者通过 /page?name=../app.py 获取应用完整源码，分析代码后可以发现更多漏洞'
        '（如SQL注入、硬编码密钥等），为后续攻击提供信息支撑。'
    )

    doc.add_paragraph('场景二：数据库窃取', style='List Number')
    doc.add_paragraph(
        '攻击者通过 /page?name=../data/users.db 下载整个数据库文件，'
        '获取所有用户的用户名、密码、邮箱、手机号等隐私数据。'
    )

    doc.add_paragraph('场景三：系统信息探测', style='List Number')
    doc.add_paragraph(
        '攻击者通过遍历系统敏感文件（如 /etc/passwd、/etc/hosts、/proc/self/environ 等）'
        '获取服务器环境信息，为服务器入侵做准备。'
    )

    doc.add_paragraph('场景四：日志文件分析', style='List Number')
    doc.add_paragraph(
        '攻击者读取应用日志文件，获取其他用户的请求记录、Session信息等敏感数据。'
    )

    # ==========================================
    # 三、漏洞成因分析
    # ==========================================
    h1 = doc.add_heading('三、漏洞成因分析', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph('该文件包含漏洞产生的根本原因包括以下几个方面：')

    causes = [
        ('1. 信任用户输入',
         '开发者将用户传入的 name 参数直接用于文件路径构造，未遵循"所有用户输入都是不可信的"这一安全原则。'),
        ('2. 缺乏路径白名单校验',
         '应用未对可访问的文件列表进行白名单限制，允许访问任意文件。正确的做法是只允许访问预定义的安全文件列表。'),
        ('3. os.path.join 不是安全函数',
         'os.path.join 仅负责拼接路径，不会阻止路径遍历攻击。攻击者可以通过 "../" 穿越出目标目录。'
         '应使用 os.path.realpath() 或 os.path.abspath() 解析后与基准目录进行比较验证。'),
        ('4. 缺乏路径规范化与验证',
         '未使用 os.path.realpath() 对最终路径进行规范化处理，也未验证解析后的路径是否在允许的目录范围内。'),
        ('5. 无文件类型限制',
         '未限制可读取的文件扩展名或MIME类型，导致攻击者可以读取任意类型文件（包括二进制数据库文件）。'),
    ]
    for title, desc in causes:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        doc.add_paragraph(desc)

    # ==========================================
    # 四、漏洞风险评估
    # ==========================================
    h1 = doc.add_heading('四、漏洞风险评估', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # 风险表格
    table = doc.add_table(rows=8, cols=2, style='Table Grid')
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(12)

    risk_data = [
        ('评估维度', '评估结果'),
        ('漏洞名称', '本地文件包含（LFI）/ 路径遍历（Path Traversal）'),
        ('CWE编号', 'CWE-22: Improper Limitation of a Pathname to a Restricted Directory'),
        ('CVSS评分', '7.5（高危 High）'),
        ('攻击向量', '网络远程（Network）'),
        ('攻击复杂度', '低（Low）— 无需任何认证即可利用'),
        ('影响范围', '机密性：高 / 完整性：无 / 可用性：无'),
        ('业务影响', '源代码泄露、数据库泄露、用户隐私泄露、服务器信息泄露'),
    ]

    for i, (col1, col2) in enumerate(risk_data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '1A3C6E')
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True

    doc.add_paragraph()

    doc.add_paragraph(
        '该漏洞无需任何认证即可利用，攻击复杂度极低，且可能造成严重的数据泄露，'
        '综合评定为高危漏洞，建议立即修复。'
    )

    # ==========================================
    # 五、漏洞修复方案
    # ==========================================
    h1 = doc.add_heading('五、漏洞修复方案', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph(
        '针对该文件包含漏洞，提供以下修复方案。方案一为推荐方案（白名单机制），'
        '方案二为补充加固方案（路径规范化验证）。建议同时采用两种方案，实现纵深防御。'
    )

    # 方案一
    h2 = doc.add_heading('方案一：白名单机制（推荐）', level=2)

    doc.add_paragraph(
        '建立允许访问的文件白名单，仅允许用户访问预定义的安全文件列表。'
        '用户的输入仅作为白名单的键（key），而非实际的文件路径。'
    )

    doc.add_paragraph('修复思路：')
    doc.add_paragraph(
        '1. 定义一个字典，将合法的页面名称映射到实际的文件路径；'
        '2. 从用户输入获取 name 参数后，先查找白名单；'
        '3. 如果 name 不在白名单中，直接返回错误提示；'
        '4. 如果在白名单中，使用预定义的安全路径读取文件。'
    )

    doc.add_paragraph('修复代码示例：')

    add_code_block(doc, '''# 定义文件白名单：用户输入 → 安全文件路径
ALLOWED_PAGES = {
    "help": "pages/help.html",
    "about": "pages/about.html",
    "contact": "pages/contact.html",
}

@app.route("/page")
def page():
    name = request.args.get("name", "")
    page_content = None

    if name:
        # 白名单校验
        if name in ALLOWED_PAGES:
            filepath = ALLOWED_PAGES[name]
            try:
                with open(filepath, encoding="utf-8") as f:
                    page_content = f.read()
            except FileNotFoundError:
                page_content = "页面文件不存在"
        else:
            page_content = "请求的页面不存在"

    return render_template("index.html", page_content=page_content)''')

    doc.add_paragraph('方案一优点：')
    advantages = [
        '从根本上杜绝路径遍历攻击，用户输入完全不参与路径构造；',
        '实现简单，易于维护；',
        '新增页面只需在 ALLOWED_PAGES 字典中添加一条记录。',
    ]
    for item in advantages:
        doc.add_paragraph(item, style='List Bullet')

    # 方案二
    h2 = doc.add_heading('方案二：路径规范化与目录限制验证', level=2)

    doc.add_paragraph(
        '如果需要更灵活的文件访问（如动态页面），可以采用路径规范化验证方案。'
        '通过对用户输入的路径进行规范化处理，并验证其是否在允许的目录范围内。'
    )

    doc.add_paragraph('修复思路：')
    doc.add_paragraph(
        '1. 定义安全的基础目录（如 pages/ 的绝对路径）；'
        '2. 将用户输入与基础目录拼接得到完整路径；'
        '3. 使用 os.path.realpath() 或 os.path.abspath() 解析为规范化的绝对路径；'
        '4. 验证规范化后的路径是否以基础目录开头（即确保文件在允许范围内）；'
        '5. 可选：限制允许的文件扩展名。'
    )

    doc.add_paragraph('修复代码示例：')

    add_code_block(doc, '''import os

# 安全基础目录
SAFE_BASE_DIR = os.path.realpath("pages")

@app.route("/page")
def page():
    name = request.args.get("name", "")
    page_content = None

    if name:
        # 过滤路径遍历字符（第一层防御）
        if ".." in name or name.startswith("/") or name.startswith("\\\\"):
            page_content = "非法的页面请求"
            return render_template("index.html", page_content=page_content)

        # 构造完整路径
        filepath = os.path.join(SAFE_BASE_DIR, name)

        # 规范化路径，消除 ../ 等符号
        real_path = os.path.realpath(filepath)

        # 验证路径是否在安全目录内
        if not real_path.startswith(SAFE_BASE_DIR):
            page_content = "非法的页面请求"
            return render_template("index.html", page_content=page_content)

        # 可选：限制文件扩展名
        allowed_ext = ('.html', '.htm', '.txt')
        if not real_path.lower().endswith(allowed_ext):
            page_content = "不支持的文件类型"
            return render_template("index.html", page_content=page_content)

        try:
            with open(real_path, encoding="utf-8") as f:
                page_content = f.read()
        except FileNotFoundError:
            page_content = "页面不存在"

    return render_template("index.html", page_content=page_content)''')

    doc.add_paragraph('方案二要点说明：')
    points = [
        'realpath() 会解析所有符号链接和相对路径（如 ../），返回规范化的绝对路径；',
        'startswith() 验证确保解析后的路径确实在允许的目录内；',
        '即使攻击者传入 ../../../etc/passwd，经 realpath() 解析后路径将不满足 startswith(SAFE_BASE_DIR) 条件而被拒绝；',
        '文件扩展名白名单进一步限制可读取的文件类型。',
    ]
    for item in points:
        doc.add_paragraph(item, style='List Bullet')

    # 方案三
    h2 = doc.add_heading('方案三：输入过滤（辅助防御）', level=2)

    doc.add_paragraph(
        '在方案一或方案二的基础上，可增加输入过滤作为辅助防御层：'
    )

    filters = [
        '过滤或拒绝包含 "../"、"..\\" 等路径遍历字符的输入；',
        '过滤或拒绝包含空字节（%00）的输入，防止截断攻击；',
        '过滤或拒绝以 "/" 或 "\\" 开头的绝对路径输入；',
        '对输入进行URL解码后再进行安全检查，防止编码绕过。',
    ]
    for item in filters:
        doc.add_paragraph(item, style='List Bullet')

    # 推荐修复方案
    h2 = doc.add_heading('综合推荐修复方案', level=2)

    doc.add_paragraph(
        '强烈推荐采用"方案一（白名单机制）+ 方案三（输入过滤）"的组合方案。原因如下：'
    )
    reasons = [
        '白名单机制从根源上切断了用户输入与文件路径的联系，安全性最高；',
        '输入过滤作为第一道防线，可提前拦截明显的恶意请求；',
        '方案组合形成纵深防御，即使一层防御被绕过，仍有另一层保护；',
        '实现和维护成本较低。',
    ]
    for item in reasons:
        doc.add_paragraph(item, style='List Bullet')

    # ==========================================
    # 六、修复验证
    # ==========================================
    h1 = doc.add_heading('六、修复验证', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph(
        '修复完成后，建议进行以下验证测试，确保漏洞已被正确修复：'
    )

    verify_tests = [
        ('正常功能验证',
         '访问 /page?name=help，应正常返回帮助页面内容。'),
        ('路径遍历验证',
         '访问 /page?name=../app.py，应返回"页面不存在"或"非法请求"，不应返回源码。'),
        ('多层遍历验证',
         '访问 /page?name=../../../etc/passwd，应返回错误提示，不应返回系统文件内容。'),
        ('数据库访问验证',
         '访问 /page?name=../data/users.db，应返回错误提示，不应返回数据库内容。'),
        ('编码绕过验证',
         '访问 /page?name=..%2Fapp.py 或 /page?name=%2E%2E%2Fapp.py，应返回错误提示。'),
        ('绝对路径验证',
         '访问 /page?name=/etc/passwd，应返回错误提示。'),
        ('空字节验证',
         '访问 /page?name=help%00.html（如适用），应返回错误提示。'),
        ('白名单边界验证',
         '访问白名单中不存在的页面名称（如 /page?name=nonexist），应返回"页面不存在"。'),
    ]

    for i, (test_name, test_desc) in enumerate(verify_tests, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {test_name}')
        run.font.bold = True
        doc.add_paragraph(f'   测试方法：{test_desc}')

    doc.add_paragraph()
    doc.add_paragraph(
        '所有验证项通过后，可确认文件包含漏洞已被成功修复。'
    )

    # ==========================================
    # 七、安全建议
    # ==========================================
    h1 = doc.add_heading('七、安全建议', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph(
        '除了修复本次发现的文件包含漏洞外，建议采取以下安全措施：'
    )

    suggestions = [
        ('最小权限原则',
         'Web应用进程应以最低权限运行，限制其对文件系统的访问范围。'
         '例如，Web用户仅能读取 Web 目录下的文件。'),
        ('禁用不安全函数',
         '审查代码中所有涉及文件操作的函数（open、os.system、subprocess、eval、exec等），'
         '确保用户输入不直接参与这些函数的参数构造。'),
        ('统一安全过滤',
         '建议实现统一的输入过滤中间件或装饰器，对所有用户输入进行安全过滤，'
         '避免在每个路由中重复编写过滤逻辑。'),
        ('定期安全审计',
         '建立定期的代码安全审计机制，将安全审查纳入开发流程（如代码审查、静态分析工具扫描）。'),
        ('安全编码培训',
         '对开发团队进行安全编码培训，重点覆盖 OWASP Top 10 中常见的Web安全漏洞类型。'),
        ('部署WAF',
         '在应用前部署Web应用防火墙（WAF），可有效拦截常见的路径遍历、SQL注入等攻击请求。'),
        ('错误信息处理',
         '生产环境应关闭调试模式（debug=False），使用自定义的错误页面，'
         '避免在错误信息中泄露文件路径、代码片段等敏感信息。'),
        ('日志监控告警',
         '对包含 "../" 等可疑特征的请求进行日志记录和告警，及时发现和响应攻击行为。'),
    ]

    for i, (title, desc) in enumerate(suggestions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title}')
        run.font.bold = True
        doc.add_paragraph(f'   {desc}')

    doc.add_page_break()

    # ==========================================
    # 附录
    # ==========================================
    h1 = doc.add_heading('附录：相关代码片段', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph('附录A：漏洞所在的原始代码（app.py 第69-91行）')
    add_code_block(doc, '''# ===== 存在文件包含漏洞的原始代码 =====

@app.route("/page")
def page():
    name = request.args.get("name", "")   # 未过滤的用户输入
    page_content = None
    if name:
        filepath = os.path.join("pages", name)   # 不安全路径拼接
        if os.path.exists(filepath):
            with open(filepath, encoding="utf-8") as f:
                page_content = f.read()           # 读取任意文件
        else:
            filepath = os.path.join("pages", name + ".html")
            if os.path.exists(filepath):
                with open(filepath, encoding="utf-8") as f:
                    page_content = f.read()
            else:
                page_content = "页面不存在"

    username = session.get("username")
    user_info = None
    if username and username in USERS:
        user_info = USERS[username]
    return render_template("index.html", username=username,
                           user=user_info, page_content=page_content)''')

    doc.add_paragraph()
    doc.add_paragraph('附录B：模板中的输出渲染（index.html 第7行）')
    add_code_block(doc, '''<!-- 模板中 page_content 使用 safe 过滤器输出，未转义 -->
{% if page_content %}
    <div class="card page-content-area">
        {{ page_content | safe }}
        <a href="/" class="btn btn-secondary">返回首页</a>
    </div>
{% endif %}''')

    doc.add_paragraph()
    doc.add_paragraph('附录C：攻击载荷示例')
    add_code_block(doc, '''# 读取应用源码
http://localhost:5000/page?name=../app.py

# 读取数据库
http://localhost:5000/page?name=../data/users.db

# 读取登录页面模板
http://localhost:5000/page?name=../templates/login.html

# 读取系统文件（Linux）
http://localhost:5000/page?name=../../../etc/passwd
http://localhost:5000/page?name=../../../etc/hosts

# URL编码绕过
http://localhost:5000/page?name=..%2Fapp.py
http://localhost:5000/page?name=%2E%2E%2Fapp.py

# 多层路径遍历
http://localhost:5000/page?name=....//....//....//etc/passwd''')

    # ========== 保存文档 ==========
    output_path = '文件包含漏洞发现与修复报告.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')


if __name__ == '__main__':
    generate_report()
