"""
生成越权与逻辑漏洞分析报告 Word 文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """设置单元格背景颜色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # 添加背景色
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)
    return p


def add_bullet(doc, text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    return p


def add_heading_styled(doc, text, level):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    return heading


def create_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r % 2 == 1:
                set_cell_shading(cell, 'D9E2F3')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后空行
    return table


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============================================================
    # 封面
    # ============================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Web 应用安全漏洞分析报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('越权漏洞与业务逻辑漏洞')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

    doc.add_paragraph()
    doc.add_paragraph()

    # 封面信息
    info_items = [
        ('项目名称', '用户管理系统（Flask Web）'),
        ('报告类型', '安全漏洞分析报告'),
        ('漏洞类型', '越权访问 / 业务逻辑漏洞'),
        ('风险等级', '高危'),
        ('报告日期', datetime.date.today().strftime('%Y年%m月%d日')),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ============================================================
    # 目录
    # ============================================================
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 概述',
        '   1.1 项目简介',
        '   1.2 测试范围',
        '2. 漏洞一：越权查看其他用户信息（高危）',
        '   2.1 漏洞描述',
        '   2.2 漏洞位置',
        '   2.3 漏洞复现',
        '   2.4 漏洞根因分析',
        '   2.5 修复方案',
        '3. 漏洞二：充值接口业务逻辑漏洞（高危）',
        '   3.1 漏洞描述',
        '   3.2 漏洞位置',
        '   3.3 漏洞复现',
        '   3.4 漏洞根因分析',
        '   3.5 修复方案',
        '4. 补充安全风险',
        '   4.1 SQL 注入风险',
        '   4.2 敏感信息泄露',
        '5. 修复优先级建议',
        '6. 总结',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # 1. 概述
    # ============================================================
    doc.add_heading('1. 概述', level=1)

    doc.add_heading('1.1 项目简介', level=2)
    doc.add_paragraph(
        '本报告针对基于 Flask 框架开发的"用户管理系统"进行安全审计。'
        '该系统提供用户注册、登录、个人信息查看、用户搜索、文件上传、余额充值等核心功能。'
        '系统采用 Python Flask + SQLite 技术栈，使用模板渲染前端页面。'
    )

    doc.add_heading('1.2 测试范围', level=2)
    doc.add_paragraph('本次安全审计重点关注以下两类安全风险：')

    add_bullet(doc, '越权漏洞（IDOR - Insecure Direct Object Reference）：攻击者通过修改请求参数中的 user_id，未经授权访问或操作其他用户的数据。')
    add_bullet(doc, '业务逻辑漏洞：充值接口缺少必要的安全校验，导致攻击者可以篡改充值金额、为任意用户充值、甚至通过负数金额进行恶意操作。')

    doc.add_paragraph()

    # ============================================================
    # 2. 漏洞一：越权查看其他用户信息
    # ============================================================
    doc.add_heading('2. 漏洞一：越权查看其他用户信息（高危）', level=1)

    doc.add_heading('2.1 漏洞描述', level=2)
    doc.add_paragraph(
        '系统的个人信息查看接口 /profile 通过 URL 参数 user_id 来指定要查看的用户。'
        '该接口未对当前登录用户身份进行校验，攻击者只需修改 user_id 参数的值，'
        '即可查看系统中任意注册用户的敏感信息，包括用户名、邮箱、手机号码、账户余额、角色权限等。'
        '此外，即使未登录也可以直接访问该接口获取用户信息。'
    )

    doc.add_heading('2.2 漏洞位置', level=2)
    doc.add_paragraph('文件：app.py，第 152-172 行 —— /profile 路由')

    p = doc.add_paragraph()
    run = p.add_run('漏洞代码：')
    run.font.bold = True

    code = '''@app.route("/profile")
def profile():
    user_id = request.args.get("user_id", "")   # 直接从URL参数获取user_id
    user = None
    if user_id:
        conn = sqlite3.connect('data/users.db')
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT * FROM users WHERE id = ?", (user_id,))
        row = c.fetchone()
        if row:
            user = dict(row)
            user["id"] = user_id
            if user["username"] in USERS:
                user["balance"] = USERS[user["username"]]["balance"]  # 泄露余额
                user["role"] = USERS[user["username"]]["role"]        # 泄露角色
            else:
                user["balance"] = 0
                user["role"] = "user"
        conn.close()
    return render_template("profile.html", user=user)
    # 注意：此处完全没有校验 session 中的登录用户身份'''
    add_code_block(doc, code)

    doc.add_heading('2.3 漏洞复现', level=2)

    doc.add_paragraph('复现步骤：', style='List Bullet')
    add_bullet(doc, '步骤1：使用账号 alice 登录系统（或直接未登录状态）。')
    add_bullet(doc, '步骤2：访问 http://localhost:5000/profile?user_id=1，即可直接查看 admin 用户的全部敏感信息（邮箱、手机号、角色、余额等）。')
    add_bullet(doc, '步骤3：修改 user_id 为 2、3、4... 可遍历查看所有用户的详细资料。')
    add_bullet(doc, '步骤4：结合 /search 搜索功能，可批量枚举用户 ID 并逐一获取敏感数据。')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('攻击示例：')
    run.font.bold = True

    create_table(doc,
        ['攻击场景', '请求示例', '结果'],
        [
            ['查看 admin 用户', 'GET /profile?user_id=1', '获取 admin 的邮箱、手机、余额(¥99,999)、角色(admin)'],
            ['遍历查看所有用户', 'GET /profile?user_id=2', '获取 alice 及其他用户的完整个人信息'],
            ['未登录直接查看', 'GET /profile?user_id=1（无Cookie）', '仍可正常获取用户数据，无任何拦截'],
            ['搜索枚举用户ID', 'GET /search?keyword=a', '返回所有含"a"的用户ID，为批量越权提供目标'],
        ],
        col_widths=[3.5, 5.0, 7.0]
    )

    doc.add_heading('2.4 漏洞根因分析', level=2)
    doc.add_paragraph('该漏洞的根本原因包括以下三点：')

    add_bullet(doc, '缺少身份认证校验：/profile 接口没有检查用户是否已登录（未检查 session["username"]），任何未登录用户都可以直接访问。')
    add_bullet(doc, '缺少授权校验：即使已登录，接口也没有验证"当前登录用户是否有权查看请求的 user_id 对应的数据"。攻击者只需知道或猜测一个 user_id，就能查看该用户的全部信息。')
    add_bullet(doc, '直接对象引用（IDOR）：用户标识 user_id 直接暴露在 URL 参数中，且数据库查询使用该参数时未附加任何"当前用户"的过滤条件。')

    add_bullet(doc, '这是一个典型的"平行越权"（Horizontal Privilege Escalation）漏洞，属于 OWASP Top 10 中的 Broken Access Control（A01:2021）。')

    doc.add_heading('2.5 修复方案', level=2)

    p = doc.add_paragraph()
    run = p.add_run('修复方案一（推荐）：添加身份认证 + 授权校验')
    run.font.bold = True

    doc.add_paragraph('在 /profile 接口中添加登录状态检查，并将查询条件限定为当前登录用户：')

    code_fix = '''@app.route("/profile")
def profile():
    # 1. 检查是否已登录
    if "username" not in session:
        return redirect("/login")

    # 2. 只允许查看自己的信息，从 session 获取当前用户身份
    username = session["username"]
    conn = sqlite3.connect('data/users.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM users WHERE username = ?", (username,))
    row = c.fetchone()
    user = dict(row) if row else None
    conn.close()

    # 3. 不从 URL 参数获取 user_id，完全杜绝越权可能
    # 如果需要管理员查看其他用户，应单独创建 admin 专用接口并做角色校验
    if user:
        if username in USERS:
            user["balance"] = USERS[username]["balance"]
            user["role"] = USERS[username]["role"]
        else:
            user["balance"] = 0
            user["role"] = "user"
    return render_template("profile.html", user=user)'''
    add_code_block(doc, code_fix)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('修复方案二：使用装饰器统一鉴权')
    run.font.bold = True

    doc.add_paragraph('创建统一的登录验证装饰器，在所有需要保护的接口上复用：')

    code_decorator = '''from functools import wraps

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "username" not in session:
            return redirect("/login")
        return f(*args, **kwargs)
    return decorated_function

# 使用方式
@app.route("/profile")
@login_required     # 一行搞定认证
def profile():
    ...'''
    add_code_block(doc, code_decorator)

    doc.add_page_break()

    # ============================================================
    # 3. 漏洞二：充值接口业务逻辑漏洞
    # ============================================================
    doc.add_heading('3. 漏洞二：充值接口业务逻辑漏洞（高危）', level=1)

    doc.add_heading('3.1 漏洞描述', level=2)
    doc.add_paragraph(
        '系统的余额充值接口 /recharge 存在严重的业务逻辑漏洞。'
        '该接口通过 POST 请求接收 user_id 和 amount 参数，直接对指定用户的余额进行加减操作。'
        '攻击者可以：'
    )
    add_bullet(doc, '为任意用户充值任意金额（无需登录，无需支付）。')
    add_bullet(doc, '传入负数金额，恶意扣减其他用户的余额。')
    add_bullet(doc, '通过修改前端隐藏表单中的 user_id，以当前登录用户身份为他人充值。')
    add_bullet(doc, '无任何支付网关或交易确认环节，充值金额直接生效。')

    doc.add_heading('3.2 漏洞位置', level=2)
    doc.add_paragraph('文件：app.py，第 175-190 行 —— /recharge 路由')

    p = doc.add_paragraph()
    run = p.add_run('漏洞代码：')
    run.font.bold = True

    code_recharge = '''@app.route("/recharge", methods=["POST"])
def recharge():
    user_id = request.form.get("user_id", "")      # 攻击者可任意修改
    amount = float(request.form.get("amount", 0))   # 可传负数！无上限校验！
    conn = sqlite3.connect('data/users.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM users WHERE id = ?", (user_id,))
    row = c.fetchone()
    if row:
        user_dict = dict(row)
        username = user_dict["username"]
        if username in USERS:
            # 直接加余额，无任何校验、无支付环节、无交易日志
            USERS[username]["balance"] = USERS[username]["balance"] + amount
    conn.close()
    return redirect(f"/profile?user_id={user_id}")
    # 注意：
    # 1. 无登录校验 —— 未登录也能调用
    # 2. 无权限校验 —— 可以为任何用户充值
    # 3. 无数值校验 —— 可传负数扣钱
    # 4. 无支付网关 —— 充值不需要实际付款
    # 5. 无日志记录 —— 无法追溯'''
    add_code_block(doc, code_recharge)

    doc.add_heading('3.3 漏洞复现', level=2)

    doc.add_paragraph('复现场景一：为任意用户充值（无需登录，无需支付）', style='List Bullet')

    doc.add_paragraph(
        '攻击者使用 curl 或 Burp Suite 直接向 /recharge 接口发送 POST 请求，'
        '将 user_id 设为 1（admin），amount 设为 10000，即可为 admin 账户增加 ¥10,000 余额：'
    )
    add_code_block(doc, '''curl -X POST http://localhost:5000/recharge \\
  -d "user_id=1&amount=10000"
# 结果：admin 余额直接增加 ¥10,000，无需任何支付''')

    doc.add_paragraph()

    doc.add_paragraph('复现场景二：通过负数金额恶意扣减他人余额', style='List Bullet')

    doc.add_paragraph(
        '攻击者将 amount 设为负数（如 -5000），可恶意扣减任意用户的余额：'
    )
    add_code_block(doc, '''curl -X POST http://localhost:5000/recharge \\
  -d "user_id=2&amount=-5000"
# 结果：alice 余额被扣减 ¥5,000，可能变为负数''')

    doc.add_paragraph()

    doc.add_paragraph('复现场景三：通过前端隐藏域修改 user_id', style='List Bullet')

    doc.add_paragraph(
        '在 profile.html 页面中，充值表单使用隐藏域 <input type="hidden" name="user_id" value="{{ user.id }}"> 传递用户ID。'
        '攻击者可在浏览器开发者工具中修改该隐藏域的值，以当前登录用户身份为其他用户充值。'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('攻击矩阵总结：')
    run.font.bold = True

    create_table(doc,
        ['攻击场景', 'HTTP 方法', '请求体', '效果'],
        [
            ['为 admin 充值 1万', 'POST', 'user_id=1&amount=10000', 'admin 余额 +¥10,000（无支付）'],
            ['为 alice 充值 10万', 'POST', 'user_id=2&amount=100000', 'alice 余额 +¥100,000'],
            ['扣减 admin 余额', 'POST', 'user_id=1&amount=-50000', 'admin 余额 -¥50,000'],
            ['无限刷余额', 'POST（循环）', 'user_id=2&amount=99999', 'alice 余额可无限增加'],
            ['未登录直接充值', 'POST', 'user_id=1&amount=1000', '无需Cookie即可充值成功'],
        ],
        col_widths=[3.5, 2.0, 5.5, 5.0]
    )

    doc.add_heading('3.4 漏洞根因分析', level=2)
    doc.add_paragraph('该漏洞的根本原因如下：')

    add_bullet(doc, '缺少认证控制（Authentication）：接口未校验用户是否登录，任何人都可以直接调用。')
    add_bullet(doc, '缺少授权控制（Authorization）：user_id 由客户端传入且被服务端直接信任，未校验"当前用户是否有权为该 user_id 操作"。')
    add_bullet(doc, '输入校验缺失（Input Validation）：amount 参数未校验是否为合法正值，允许负数、零、超大数值。')
    add_bullet(doc, '缺少支付网关（Business Logic Gap）：余额变更没有通过任何真实的支付/交易系统，直接操作数据库。')
    add_bullet(doc, '缺少审计日志（Audit Log）：没有记录谁、在什么时间、为谁、充值了多少金额。')
    add_bullet(doc, '数据存储不一致：用户余额存储在内存字典 USERS 中而非数据库表中，服务重启后数据丢失，且并发操作不安全。')

    doc.add_heading('3.5 修复方案', level=2)

    p = doc.add_paragraph()
    run = p.add_run('修复方案一（推荐）：完整的充值接口安全加固')
    run.font.bold = True

    doc.add_paragraph('综合修复代码：')

    code_recharge_fix = '''@app.route("/recharge", methods=["POST"])
def recharge():
    # 1. 登录校验 —— 未登录无法充值
    if "username" not in session:
        flash("请先登录")
        return redirect("/login")

    # 2. 获取当前登录用户（不从客户端接受 user_id）
    username = session["username"]

    # 3. 金额校验
    try:
        amount = float(request.form.get("amount", 0))
    except (ValueError, TypeError):
        flash("金额格式错误")
        return redirect(f"/profile")

    # 金额必须为正数，且在合理范围内
    if amount <= 0:
        flash("充值金额必须大于0")
        return redirect(f"/profile")
    if amount > 50000:  # 单笔上限
        flash("单笔充值金额不能超过 ¥50,000")
        return redirect(f"/profile")

    # 4. 调用支付网关（伪代码，实际需接入支付宝/微信支付等）
    # order_id = payment_gateway.create_order(username, amount)
    # if not payment_gateway.verify_payment(order_id):
    #     flash("支付未完成")
    #     return redirect(f"/profile")

    # 5. 更新余额（应从数据库获取，非内存字典）
    conn = sqlite3.connect('data/users.db')
    c = conn.cursor()
    # 确保余额字段在数据库中存在
    c.execute(
        "UPDATE users SET balance = balance + ? WHERE username = ?",
        (amount, username)
    )
    conn.commit()

    # 6. 记录交易日志
    c.execute(
        "INSERT INTO recharge_logs (username, amount, ip, created_at) VALUES (?, ?, ?, ?)",
        (username, amount, request.remote_addr, datetime.now())
    )
    conn.commit()
    conn.close()

    flash(f"成功充值 ¥{amount:.2f}")
    return redirect(f"/profile")'''
    add_code_block(doc, code_recharge_fix)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('修复方案二：将余额存入数据库 + 使用数据库事务保证一致性')
    run.font.bold = True

    doc.add_paragraph('在数据库 users 表中添加 balance 字段，避免使用内存字典存储余额：')

    code_db_fix = '''-- 修改 users 表，添加 balance 字段
ALTER TABLE users ADD COLUMN balance REAL DEFAULT 0;

-- 创建充值日志表（用于审计）
CREATE TABLE recharge_logs (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    username TEXT NOT NULL,
    amount REAL NOT NULL,
    ip TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);'''

    add_code_block(doc, code_db_fix)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('修复方案三：CSRF 防护')
    run.font.bold = True

    doc.add_paragraph(
        '为充值等敏感操作添加 CSRF Token 验证，防止跨站请求伪造攻击。'
        '可使用 Flask-WTF 扩展或手动实现 CSRF Token 机制。'
    )

    doc.add_page_break()

    # ============================================================
    # 4. 补充安全风险
    # ============================================================
    doc.add_heading('4. 补充安全风险', level=1)

    doc.add_heading('4.1 SQL 注入风险（高危）', level=2)
    doc.add_paragraph(
        '在注册接口 /register（app.py 第101行）和搜索接口 /search（app.py 第119行）中，'
        '存在 SQL 注入漏洞，原因是使用 Python f-string 直接拼接 SQL 语句。'
    )

    doc.add_paragraph('漏洞代码示例：')

    code_sqli1 = '''# 注册接口 - SQL注入
sql = f"INSERT INTO users (username, password, email, phone) VALUES ('{username}', '{password}', '{email}', '{phone}')"
c.execute(sql)'''
    add_code_block(doc, code_sqli1)

    code_sqli2 = '''# 搜索接口 - SQL注入
sql = f"SELECT * FROM users WHERE username LIKE '%{keyword}%' OR email LIKE '%{keyword}%'"
c.execute(sql)'''
    add_code_block(doc, code_sqli2)

    doc.add_paragraph('修复方案：使用参数化查询，所有现有代码中已正确使用参数化查询的地方（如 login 和 profile 接口）应保持，不一致的地方应统一：')

    code_sqli_fix = '''# 正确做法：参数化查询
c.execute(
    "INSERT INTO users (username, password, email, phone) VALUES (?, ?, ?, ?)",
    (username, password, email, phone)
)

c.execute(
    "SELECT * FROM users WHERE username LIKE ? OR email LIKE ?",
    (f"%{keyword}%", f"%{keyword}%")
)'''
    add_code_block(doc, code_sqli_fix)

    doc.add_heading('4.2 敏感信息泄露（中危）', level=2)
    doc.add_paragraph('系统中存在多处敏感信息泄露问题：')

    add_bullet(doc, '登录页（login.html）HTML 注释中直接暴露了管理员默认账号密码：<!-- 调试信息 - 默认管理员账号 用户名: admin 密码: admin123 -->，任何访问登录页的人都能看到。')
    add_bullet(doc, '首页（index.html）直接显示用户密码明文：<li><span class="label">密码：</span>{{ user.password }}</li>。')
    add_bullet(doc, '管理员密码硬编码在 app.py 的 USERS 字典中（第23-40行），且以明文存储。')
    add_bullet(doc, 'app.py 中 Flask secret_key 硬编码为 "dev-key-2025"（第5行），且包含"dev"字样表明是开发环境密钥，生产环境使用存在 Session 伪造风险。')
    add_bullet(doc, '服务以 debug=True 模式运行：app.run(debug=True, host="0.0.0.0", port=5000)，生产环境开启 debug 模式存在 Werkzeug 调试控制台远程代码执行风险。')

    doc.add_paragraph()
    doc.add_paragraph('修复方案：')
    add_bullet(doc, '删除所有 HTML 源码中的调试注释和硬编码凭据。')
    add_bullet(doc, '页面中禁止展示密码明文，已存储的密码应使用 bcrypt/scrypt/argon2 进行哈希处理。')
    add_bullet(doc, '将 secret_key 和数据库凭据迁移至环境变量或配置文件（如 .env），并从代码仓库中排除。')
    add_bullet(doc, '生产环境必须设置 debug=False，并配合 WSGI 服务器（如 Gunicorn/uWSGI）运行。')

    doc.add_page_break()

    # ============================================================
    # 5. 修复优先级建议
    # ============================================================
    doc.add_heading('5. 修复优先级建议', level=1)

    create_table(doc,
        ['优先级', '漏洞', '风险等级', '影响', '建议修复时间'],
        [
            ['P0-紧急', '充值接口逻辑漏洞', '严重', '无限刷余额、恶意扣款，直接造成资产损失', '24小时内'],
            ['P0-紧急', '越权查看用户信息', '高危', '批量泄露用户敏感信息，违反数据保护法规', '24小时内'],
            ['P1-高', 'SQL注入漏洞', '高危', '可导致数据库被拖库、数据被篡改', '3个工作日内'],
            ['P2-中', '敏感信息泄露', '中危', '管理员凭据泄露、Session伪造风险', '1周内'],
            ['P2-中', 'Debug模式开启', '中危', '生产环境远程代码执行风险', '上线前必须修复'],
        ],
        col_widths=[2.0, 3.5, 1.8, 5.5, 3.2]
    )

    doc.add_page_break()

    # ============================================================
    # 6. 总结
    # ============================================================
    doc.add_heading('6. 总结', level=1)

    doc.add_paragraph(
        '本次安全审计共发现 2 个主要高危漏洞和多个附加安全风险：'
    )

    doc.add_paragraph(
        '越权漏洞（IDOR）方面，/profile 接口完全缺少身份认证和授权校验，'
        '攻击者可以遍历查看任意用户的敏感信息（邮箱、手机、余额、角色等）。'
        '修复核心是在接口层面添加 session 校验，并使用当前登录用户身份查询而非信任客户端传入的 user_id。'
    )

    doc.add_paragraph(
        '业务逻辑漏洞方面，/recharge 接口缺少认证、授权、输入校验、支付网关、审计日志等'
        '多层安全防护。攻击者可以为任意用户充值任意金额（含负数），直接造成资产损失。'
        '修复核心是建立完整的充值业务安全链路：认证 → 授权 → 输入校验 → 支付网关验证 → 余额变更 → 审计日志。'
    )

    doc.add_paragraph(
        '此外，建议开发团队建立以下安全开发规范：'
    )
    add_bullet(doc, '所有接口默认需要登录认证，除非明确设计为公开接口。')
    add_bullet(doc, '用户身份标识应从 Session/Token 中获取，不应信任客户端传入的参数。')
    add_bullet(doc, '所有数据库操作必须使用参数化查询，禁止字符串拼接 SQL。')
    add_bullet(doc, '敏感配置（密钥、凭据）使用环境变量管理，禁止硬编码。')
    add_bullet(doc, '生产环境关闭 Debug 模式，使用 HTTPS 传输。')
    add_bullet(doc, '引入安全测试环节，在 CI/CD 流程中加入 SAST 和 DAST 扫描。')

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'报告生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ============================================================
    # 保存文件
    # ============================================================
    output_path = r'C:\Users\吴昶豪\Desktop\越权与逻辑漏洞\越权与逻辑漏洞分析报告.docx'
    doc.save(output_path)
    print(f'报告已生成：{output_path}')


if __name__ == '__main__':
    main()
